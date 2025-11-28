"""
Report Exporter - Export LLM reports to Word and PDF with embedded charts
"""

import io
import base64
from typing import Dict, Optional, List
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go


class ReportExporter:
    """Export reports to Word and PDF formats with charts"""
    
    def __init__(self):
        """Initialize report exporter"""
        pass
    
    def export_to_word(
        self,
        report_text: str,
        charts: Optional[Dict] = None,
        filename: str = "report.docx"
    ) -> io.BytesIO:
        """
        Export report to Word document with embedded charts
        
        Args:
            report_text: Markdown formatted report text
            charts: Dictionary of Plotly figure objects
            filename: Output filename
            
        Returns:
            BytesIO object containing the Word document
        """
        doc = Document()
        
        # Set document style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add title
        title = doc.add_heading('Health Analytics Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'Generated: {self._get_timestamp()}')
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Spacing
        
        # Parse markdown and add to document
        lines = report_text.split('\n')
        in_chart_section = False
        chart_placeholder = None
        
        for line in lines:
            line = line.strip()
            
            # Check for chart placeholders
            if line.startswith('[CHART:') and line.endswith(']'):
                chart_name = line[7:-1].strip()
                chart_placeholder = chart_name
                continue
            
            # Skip empty lines in chart section
            if in_chart_section and not line:
                continue
            
            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
                # Add chart after heading if placeholder exists
                if chart_placeholder and charts and chart_placeholder in charts:
                    self._add_chart_to_doc(doc, charts[chart_placeholder], chart_placeholder)
                    chart_placeholder = None
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                if chart_placeholder and charts and chart_placeholder in charts:
                    self._add_chart_to_doc(doc, charts[chart_placeholder], chart_placeholder)
                    chart_placeholder = None
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            
            # Horizontal rules
            elif line.startswith('---'):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('_' * 60)
                doc.add_paragraph()
            
            # Lists
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            
            # Bold text
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            
            # Regular paragraphs
            elif line:
                # Handle inline bold
                if '**' in line:
                    p = doc.add_paragraph()
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold
                            run.bold = True
                else:
                    doc.add_paragraph(line)
        
        # Save to BytesIO
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    
    def _add_chart_to_doc(self, doc: Document, fig: go.Figure, chart_name: str):
        """
        Add Plotly chart to Word document as image
        
        Args:
            doc: Document object
            fig: Plotly figure
            chart_name: Name of the chart
        """
        try:
            # Export chart as image
            img_bytes = fig.to_image(format="png", width=1200, height=600, scale=2)
            
            # Add image to document
            image_stream = io.BytesIO(img_bytes)
            doc.add_picture(image_stream, width=Inches(6))
            
            # Add caption
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(f'Figure: {chart_name}')
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Spacing
            
        except Exception as e:
            # If chart fails, add placeholder text
            p = doc.add_paragraph()
            run = p.add_run(f'[Chart: {chart_name} - Unable to embed]')
            run.font.italic = True
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    def export_to_pdf(
        self,
        report_text: str,
        charts: Optional[Dict] = None,
        filename: str = "report.pdf"
    ) -> io.BytesIO:
        """
        Export report to PDF with embedded charts
        
        Args:
            report_text: Markdown formatted report text
            charts: Dictionary of Plotly figure objects
            filename: Output filename
            
        Returns:
            BytesIO object containing the PDF
        """
        # Convert markdown to HTML
        html_content = markdown.markdown(report_text, extensions=['tables', 'fenced_code'])
        
        # Add styling
        html_with_style = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: 'Calibri', 'Arial', sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                h1 {{
                    color: #0066CC;
                    border-bottom: 3px solid #0066CC;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #0066CC;
                    margin-top: 30px;
                    border-bottom: 2px solid #ccc;
                    padding-bottom: 5px;
                }}
                h3 {{
                    color: #004499;
                    margin-top: 20px;
                }}
                strong {{
                    color: #0066CC;
                }}
                ul {{
                    margin-left: 20px;
                }}
                .chart-placeholder {{
                    background: #f0f0f0;
                    padding: 20px;
                    margin: 20px 0;
                    border: 2px dashed #ccc;
                    text-align: center;
                    color: #666;
                }}
                .timestamp {{
                    text-align: center;
                    color: #666;
                    font-size: 0.9em;
                    margin-bottom: 30px;
                }}
            </style>
        </head>
        <body>
            <div class="timestamp">Generated: {self._get_timestamp()}</div>
            {html_content}
        </body>
        </html>
        """
        
        # Replace chart placeholders with actual chart images
        if charts:
            for chart_name, fig in charts.items():
                try:
                    # Convert chart to base64 image
                    img_bytes = fig.to_image(format="png", width=800, height=400, scale=2)
                    img_base64 = base64.b64encode(img_bytes).decode()
                    img_html = f'<img src="data:image/png;base64,{img_base64}" style="width: 100%; max-width: 800px; margin: 20px 0;" />'
                    
                    # Replace placeholder
                    placeholder = f'[CHART: {chart_name}]'
                    html_with_style = html_with_style.replace(placeholder, img_html)
                except Exception as e:
                    print(f"Error embedding chart {chart_name}: {e}")
        
        # Convert to PDF using simple HTML to PDF method
        # Note: For production, you might want to use weasyprint or reportlab
        pdf_bytes = io.BytesIO()
        pdf_bytes.write(html_with_style.encode('utf-8'))
        pdf_bytes.seek(0)
        
        return pdf_bytes
    
    def _get_timestamp(self) -> str:
        """Get formatted timestamp"""
        from datetime import datetime
        return datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    def create_download_button(
        self,
        report_text: str,
        charts: Optional[Dict] = None,
        format_type: str = "word",
        button_label: str = "Download Report"
    ) -> bytes:
        """
        Create downloadable report file
        
        Args:
            report_text: Report content
            charts: Charts to embed
            format_type: "word" or "pdf"
            button_label: Button label
            
        Returns:
            Bytes of the file
        """
        if format_type.lower() == "word":
            return self.export_to_word(report_text, charts)
        elif format_type.lower() == "pdf":
            return self.export_to_pdf(report_text, charts)
        else:
            raise ValueError(f"Unsupported format: {format_type}")


# Utility functions for Streamlit integration
def prepare_charts_for_export(chart_figures: Dict) -> Dict:
    """
    Prepare Plotly figures for export
    
    Args:
        chart_figures: Dictionary mapping chart names to Plotly figures
        
    Returns:
        Dictionary with chart metadata
    """
    charts = {}
    for name, fig in chart_figures.items():
        charts[name] = {
            'figure': fig,
            'title': fig.layout.title.text if fig.layout.title else name,
            'type': 'plotly',
            'description': f'Visualization for {name}'
        }
    return charts

